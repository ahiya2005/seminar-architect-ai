import streamlit as st
import google.generativeai as genai
import io
import time
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import PyPDF2

# --- 1. הגדרות דף ---
st.set_page_config(page_title="Seminar Architect PRO", page_icon="🎓", layout="wide")

api_key = st.secrets.get("GEMINI_API_KEY")
if api_key:
    genai.configure(api_key=api_key)

st.markdown("""
    <style>
    header {visibility: hidden;}
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    div[data-testid="InputInstructions"] { display: none !important; }
    .stButton>button { width: 100%; background-color: #2c3e50; color: white; border-radius: 10px; font-weight: bold; height: 3.5em; }
    .stProgress > div > div > div { background-image: linear-gradient(45deg, rgba(255, 255, 255, .15) 25%, transparent 25%, transparent 50%, rgba(255, 255, 255, .15) 50%, rgba(255, 255, 255, .15) 75%, transparent 75%, transparent); background-size: 1rem 1rem; animation: progress-bar-stripes 1s linear infinite; }
    @keyframes progress-bar-stripes { from { background-position: 1rem 0; } to { background-position: 0 0; } }
    </style>
""", unsafe_allow_html=True)

# --- 2. פונקציות עזר (ניקוי ו-RTL) ---
def is_valid_email(email):
    return re.match(r"^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z]{2,}$", email) is not None

def aggressive_clean_tags(text):
    text = re.sub(r"\[.*?source.*?\]", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\[\d+\]", "", text)
    return text.strip()

def extract_text_from_file(uploaded_file):
    if uploaded_file is None: return ""
    try:
        if uploaded_file.name.endswith('.pdf'):
            return "\n".join([page.extract_text() for page in PyPDF2.PdfReader(uploaded_file).pages])
        return uploaded_file.getvalue().decode("utf-8")
    except Exception: return ""

def set_rtl_paragraph(p):
    pPr = p._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def add_rtl_run(paragraph, text, font_name='David', font_size=12, bold=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    return run

# --- 3. הליבה: מנוע ה-AI היציב (ללא חריגת בקשות) ---
def call_gemini_safe(prompt, max_tokens=8192):
    model = genai.GenerativeModel('gemini-1.5-flash')
    error_msg = "Unknown Error"
    
    for attempt in range(3):
        try:
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(temperature=0.4, max_output_tokens=max_tokens),
                safety_settings=[
                    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
                ]
            )
            try:
                if response.text: return aggressive_clean_tags(response.text)
            except ValueError:
                # קורה אם גוגל חסמה לחלוטין את הטקסט בגלל צנזורה
                return "שגיאה: התוכן נחסם על ידי מסנני הבטיחות של גוגל."
                
        except Exception as e:
            error_msg = str(e)
            if "429" in error_msg or "Quota" in error_msg:
                time.sleep(30) # המתנה משמעותית לאיפוס מכסה
            else:
                time.sleep(5)
                
    return f"שגיאה בהתקשרות מול גוגל: {error_msg}"

def get_dynamic_outline(topic, extra):
    prompt = f"תפקיד: פרופסור. צור 6 כותרות אקדמיות לפרקים בנושא '{topic}'. הנחיות: {extra}. פרק אחרון: 'ביבליוגרפיה'. החזר רק את הכותרות מופרדות בפסיק (без מספור)."
    res = call_gemini_safe(prompt, 500)
    if res and "שגיאה" not in res: 
        return [c.strip() for c in res.split(',') if c.strip()][:6]
    return ["מבוא", "רקע תיאורטי", "ניתוח מערכות", "מקרי בוחן", "מסקנות", "ביבליוגרפיה"]

def generate_chapter(chapter_title, topic, name, extra, notes):
    if "ביבליוגרפיה" in chapter_title or "מקורות" in chapter_title:
        prompt = f"תפקיד: ביבליוגרף. צור רשימת מקורות אקדמית (APA) בנושא '{topic}'. החזר רשימה בלבד (12 מקורות) ללא פסקאות הסבר."
        content = call_gemini_safe(prompt, 1500)
        return f"# {chapter_title}\n{content}"
        
    prompt = f"""
    תפקיד: פרופסור אקדמי בכיר.
    משימה: כתוב פרק עומק אקדמי תחת הכותרת '{chapter_title}' לעבודה בנושא '{topic}' עבור הסטודנט {name}.
    הנחיות: {extra}. סילבוס: {notes}.
    
    חוקי ברזל חובה:
    1. אורך ועומק: הפרק חייב להיות ארוך ומעמיק מאוד (לפחות 800 מילים).
    2. מבנה פנימי: עליך לחלק את הפרק ל-3 תתי-נושאים. תן לכל תת-נושא כותרת עם הסימן '##'.
    3. תחת כל תת-נושא (##) כתוב לפחות 3 פסקאות ארוכות ובשרניות.
    4. ציטוטים: חובה לשלב ציטוטי APA פנימיים (מחבר, שנה) כדי לגבות כל טענה.
    5. תגיות: איסור מוחלט על שימוש בסוגריים מרובעים כגון .
    6. ללא הקדמות: התחל מיד עם הכותרת '# {chapter_title}'.
    """
    content = call_gemini_safe(prompt, 8192)
    return content

# --- 4. בניית הוורד המקצועי ---
def create_master_doc(topic, author, institution, content_list, lang):
    doc = Document()
    font_name = 'David' if lang == "עברית" else 'Times New Roman'
    
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.98)

    # שער מדויק
    doc.add_paragraph('\n\n\n\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl_paragraph(p)
    add_rtl_run(p, f"עבודה סמינריונית בנושא:\n{topic}", font_name, 24, True)
    
    doc.add_paragraph('\n\n\n\n')
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl_paragraph(p2)
    add_rtl_run(p2, f"מוגש על ידי: {author}", font_name, 16, False)
    if institution:
        add_rtl_run(p2, "\n", font_name, 16, False)
        add_rtl_run(p2, f"מוסד אקדמי: {institution}", font_name, 16, False)
    doc.add_page_break()

    # טקסט
    for text in content_list:
        for line in text.split('\n'):
            line = line.strip()
            if not line: continue
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl_paragraph(p)
            
            if line.startswith('# ') and not line.startswith('## '):
                add_rtl_run(p, line.replace('#', '').strip(), font_name, 18, True)
            elif line.startswith('## '):
                add_rtl_run(p, line.replace('##', '').strip(), font_name, 14, True)
            else:
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.alignment = 3 
                add_rtl_run(p, line.replace('*', ''), font_name, 12, False)
                
        doc.add_page_break()
    return doc

# --- 5. ממשק משתמש ---
lang = st.radio("🌐 שפת ממשק:", ["עברית", "English"], horizontal=True)
if lang == "עברית":
    st.markdown("<style>.block-container { direction: rtl; text-align: right; }</style>", unsafe_allow_html=True)

if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False

st.title("🎓 Seminar Architect PRO")

if not st.session_state['logged_in']:
    email_input = st.text_input("אימייל להתחברות:")
    if st.button("🚀 כניסה למערכת"):
        if is_valid_email(email_input):
            st.session_state['user_email'] = email_input.lower()
            st.session_state['logged_in'] = True
            st.rerun()
        else: st.error("נא להזין אימייל תקין.")
else:
    st.sidebar.success(f"מחובר כ: {st.session_state['user_email']}")
    col1, col2 = st.columns(2)
    with col1:
        topic = st.text_input("נושא הסמינריון:")
        name = st.text_input("שם הסטודנט:")
        institution = st.text_input("מוסד אקדמי (אופציונלי):")
    with col2:
        extra = st.text_area("דגשים ספציפיים ופרוטוקול אישי:", height=130)
    uploaded = st.file_uploader("העלאת סילבוס (PDF/TXT):", type=['pdf', 'txt'])

    if st.button("🚀 צא לדרך! הפעל בנייה חכמה ובטוחה"):
        if not topic or not name: st.error("הזן נושא ושם סטודנט!")
        elif not api_key: st.error("שגיאת API KEY.")
        else:
            st.warning("⚠️ המערכת בונה את העבודה בקצב מבוקר כדי למנוע חסימות גוגל (כ-2 דקות). נא לא לסגור!")
            notes = extract_text_from_file(uploaded)
            
            status_text = st.empty()
            progress_bar = st.progress(0)
            
            status_text.info("⏳ מאבחן נושא ובונה שלד פרקים...")
            chapters = get_dynamic_outline(topic, extra)
            
            generated_content = []
            total_steps = len(chapters) + 1 
            
            for i, head in enumerate(chapters):
                status_text.info(f"⏳ כותב פרק עומק: **{head}**...")
                content = generate_chapter(head, topic, name, extra, notes)
                generated_content.append(content)
                progress_bar.progress((i + 1) / total_steps)
                
                # מנגנון ביטחון: מנוחה של 10 שניות בין פרק לפרק מונעת את שגיאת ה-429 (עומס) לחלוטין
                time.sleep(10)
            
            status_text.info("⏳ מסכם את העבודה ומייצר תקציר...")
            abs_prompt = f"תפקיד: פרופסור. כתוב תקציר אקדמי (כ-300 מילים) לעבודה בנושא '{topic}'. התחל ישירות עם '# תקציר'."
            abstract = call_gemini_safe(abs_prompt, 1000)
            generated_content.insert(0, abstract) 
            
            progress_bar.progress(1.0)
            status_text.success("🎉 הסמינריון הושלם בהצלחה מרובה!")
            
            doc = create_master_doc(topic, name, institution, generated_content, lang)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button("📥 הורד סמינריון מושלם (Word)", buf, f"Seminar_{name}.docx")
            st.balloons()
